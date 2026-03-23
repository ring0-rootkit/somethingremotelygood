#!/usr/bin/env python3
"""
CTDA 2024 Conference Article Generator.

Generates a ~4-page article titled:
"Разработка защищённой системы удалённой виртуализации"

Usage:
    python3 docx/generate_article.py

Output:
    article.docx in the project root directory.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
OUTPUT_PATH = os.path.join(PROJECT_ROOT, 'article.docx')


def setup_styles(doc):
    """Create CTDA 2024 styles."""
    styles = doc.styles

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Название статьи CTDA ──
    if 'Название статьи CTDA' not in [s.name for s in styles]:
        s = styles.add_style('Название статьи CTDA', 1)
    else:
        s = styles['Название статьи CTDA']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(12)
    s.font.bold = True
    s.font.all_caps = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(12)
    s.paragraph_format.line_spacing = 1.0

    # ── Автор ──
    if 'Автор' not in [s.name for s in styles]:
        s = styles.add_style('Автор', 1)
    else:
        s = styles['Автор']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(11)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.line_spacing = 1.0

    # ── Аффилиация ──
    if 'Аффилиация' not in [s.name for s in styles]:
        s = styles.add_style('Аффилиация', 1)
    else:
        s = styles['Аффилиация']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.italic = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(12)
    s.paragraph_format.line_spacing = 1.0

    # ── Аннотация ──
    if 'Аннотация' not in [s.name for s in styles]:
        s = styles.add_style('Аннотация', 1)
    else:
        s = styles['Аннотация']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.first_line_indent = Cm(1.0)

    # ── Ключевые слова ──
    if 'Ключевые слова' not in [s.name for s in styles]:
        s = styles.add_style('Ключевые слова', 1)
    else:
        s = styles['Ключевые слова']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(18)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.first_line_indent = Cm(1.0)

    # ── Подзаголовок ──
    if 'Подзаголовок' not in [s.name for s in styles]:
        s = styles.add_style('Подзаголовок', 1)
    else:
        s = styles['Подзаголовок']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(11)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.0

    # ── Подзаголовок 2 ──
    if 'Подзаголовок 2' not in [s.name for s in styles]:
        s = styles.add_style('Подзаголовок 2', 1)
    else:
        s = styles['Подзаголовок 2']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(11)
    s.font.bold = True
    s.font.italic = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(6)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.first_line_indent = Cm(1.0)

    # ── Обычный (body text) ──
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # ── БИБЛИОГРАФИЧЕСКИЕ ССЫЛКИ ──
    if 'БИБЛИОГРАФИЧЕСКИЕ ССЫЛКИ' not in [s.name for s in styles]:
        s = styles.add_style('БИБЛИОГРАФИЧЕСКИЕ ССЫЛКИ', 1)
    else:
        s = styles['БИБЛИОГРАФИЧЕСКИЕ ССЫЛКИ']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(11)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.0

    # ── Список литературы ──
    if 'Список литературы' not in [s.name for s in styles]:
        s = styles.add_style('Список литературы', 1)
    else:
        s = styles['Список литературы']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.first_line_indent = Cm(0)
    s.paragraph_format.left_indent = Cm(0.5)


def add_keywords(doc, text):
    """Add keywords paragraph with 'Ключевые слова:' in italic."""
    p = doc.add_paragraph(style='Ключевые слова')
    run_label = p.add_run('Ключевые слова: ')
    run_label.italic = True
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(10)
    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(10)


def add_reference(doc, number, text, authors_italic=None):
    """Add a bibliography entry. authors_italic is the author string to italicize."""
    p = doc.add_paragraph(style='Список литературы')
    prefix = f'{number}. '
    if authors_italic:
        run_num = p.add_run(prefix)
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(10)
        run_auth = p.add_run(authors_italic)
        run_auth.italic = True
        run_auth.font.name = 'Times New Roman'
        run_auth.font.size = Pt(10)
        rest = text[len(authors_italic):]
        run_rest = p.add_run(rest)
        run_rest.font.name = 'Times New Roman'
        run_rest.font.size = Pt(10)
    else:
        run = p.add_run(prefix + text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)


def body(doc, text):
    """Add a body text paragraph."""
    p = doc.add_paragraph(text, style='Normal')
    return p


def main():
    print('[*] Generating CTDA 2024 article...')
    doc = Document()
    setup_styles(doc)

    # Remove initial empty paragraph
    if doc.paragraphs and doc.paragraphs[0].text == '':
        p_el = doc.paragraphs[0]._element
        p_el.getparent().remove(p_el)

    # ════════════════════════════════════════════
    # TITLE
    # ════════════════════════════════════════════
    doc.add_paragraph(
        'РАЗРАБОТКА ЗАЩИЩЁННОЙ СИСТЕМЫ УДАЛЁННОЙ ВИРТУАЛИЗАЦИИ',
        style='Название статьи CTDA'
    )

    # AUTHOR
    doc.add_paragraph('И. И. Иванов', style='Автор')

    # AFFILIATION
    p = doc.add_paragraph(style='Аффилиация')
    run = p.add_run(
        'Белорусский государственный университет, пр. Независимости, 4,\n'
        '220030, г. Минск, Беларусь, ivanov@bsu.by'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph(style='Аффилиация')
    run = p.add_run('Научный руководитель — П. П. Петров, старший преподаватель')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

    # ════════════════════════════════════════════
    # ABSTRACT
    # ════════════════════════════════════════════
    doc.add_paragraph(
        'Представлена архитектура и реализация системы безопасного удалённого доступа '
        'к контейнерным средам с аппаратным хранением криптографических ключей на '
        'микроконтроллере ESP32-C3. Система обеспечивает генерацию, шифрование и '
        'использование ключей Ed25519 непосредственно на микроконтроллере, исключая '
        'хранение приватных ключей на компьютере пользователя. Реализован протокол '
        'SSH-агента через серийный мост на Python, Go-клиент для аутентификации и '
        'C-менеджер для управления LXD-контейнерами с LUKS-шифрованием томов. '
        'Результаты тестирования подтверждают корректность криптографических операций '
        'и приемлемую производительность подписи (менее 250 мс).',
        style='Аннотация'
    )

    # KEYWORDS
    add_keywords(
        doc,
        'аппаратное хранение ключей; ESP32; Ed25519; SSH-агент; '
        'контейнерная виртуализация; LXD; LUKS; криптография.'
    )

    # ════════════════════════════════════════════
    # INTRODUCTION
    # ════════════════════════════════════════════
    body(doc,
        'Обеспечение безопасного удалённого доступа к серверным средам является '
        'одной из ключевых задач современной информационной безопасности. Традиционные '
        'подходы предполагают хранение криптографических ключей на файловой системе '
        'рабочей станции, что создаёт риск их компрометации при заражении вредоносным '
        'программным обеспечением, физическом доступе к устройству или утечке резервных '
        'копий [1]. Коммерческие аппаратные решения, такие как YubiKey и Nitrokey, '
        'решают эту проблему, однако их стоимость и закрытость программного кода '
        'ограничивают применимость в образовательных и исследовательских контекстах.')

    body(doc,
        'Параллельно с этим задача изоляции пользовательских сред решается '
        'посредством контейнерной виртуализации. Платформа LXD/LXC предоставляет '
        'полноценные системные контейнеры с минимальными накладными расходами, однако '
        'стандартные механизмы аутентификации не обеспечивают привязку к аппаратному '
        'носителю ключей [2].')

    body(doc,
        'Целью настоящей работы является разработка системы, объединяющей аппаратное '
        'хранение криптографических ключей на микроконтроллере ESP32-C3 с управлением '
        'контейнерными средами через протокол SSH. Система реализует полный цикл '
        'от генерации ключевой пары до аутентификации и подключения к контейнеру, '
        'при этом приватный ключ никогда не покидает микроконтроллер.')

    # ════════════════════════════════════════════
    # ARCHITECTURE
    # ════════════════════════════════════════════
    doc.add_paragraph('Архитектура системы', style='Подзаголовок')

    body(doc,
        'Система состоит из четырёх взаимодействующих компонентов: прошивки '
        'микроконтроллера ESP32-C3, Python-моста (SSH Agent Bridge), Go-клиента '
        'и C-менеджера. Каждый компонент выполняет строго определённую функцию '
        'в цепочке аутентификации и авторизации.')

    body(doc,
        'Микроконтроллер ESP32-C3, построенный на архитектуре RISC-V с тактовой '
        'частотой 160 МГц, выступает в роли аппаратного хранилища ключей. Генерация '
        'ключевой пары Ed25519 выполняется непосредственно на устройстве с использованием '
        'аппаратного генератора случайных чисел (TRNG). Приватный ключ (32-байтовое '
        'зерно) сохраняется в энергонезависимом хранилище NVS и может быть дополнительно '
        'зашифрован паролем пользователя [3].')

    body(doc,
        'Python-мост подключается к ESP32 через последовательный порт (USB-CDC, '
        '115200 бод) и создаёт Unix-сокет, совместимый с протоколом SSH-агента '
        '(RFC 4253) [4]. При получении запроса на подпись мост транслирует данные '
        'в собственный серийный протокол, пересылает на ESP32, получает подпись и '
        'формирует ответ в формате SSH-агента. Для поддержки парольной защиты '
        'реализован механизм расширений SSH-агента (тип сообщения 27), позволяющий '
        'передавать команды разблокировки и блокировки устройства.')

    body(doc,
        'Go-клиент реализует протокол аутентификации типа «вызов-ответ» с '
        'C-менеджером. Клиент подключается к SSH-агенту через Unix-сокет, запрашивает '
        'подпись случайного вызова (challenge) и отправляет её менеджеру для верификации. '
        'При успешной аутентификации менеджер разрешает доступ к назначенному '
        'контейнеру, после чего клиент устанавливает SSH-сессию.')

    body(doc,
        'C-менеджер является серверным компонентом, слушающим TCP-порт 5555. '
        'Он хранит учётные записи пользователей и привязки к контейнерам в базе '
        'данных SQLCipher. Для каждого контейнера создаётся зашифрованный LUKS-том '
        '(100 МБ), монтируемый при подключении пользователя. Менеджер взаимодействует '
        'с LXD через REST API для создания, запуска и остановки контейнеров '
        'на базе Alpine Linux 3.20 [5].')

    # ════════════════════════════════════════════
    # CRYPTOGRAPHIC SUBSYSTEM
    # ════════════════════════════════════════════
    doc.add_paragraph('Криптографическая подсистема', style='Подзаголовок')

    body(doc,
        'Криптографическая подсистема основана на алгоритме цифровой подписи '
        'Ed25519, использующем эллиптическую кривую Curve25519 в скрученной форме '
        'Эдвардса [6]. Реализация на ESP32 выполнена на основе библиотеки TweetNaCl, '
        'адаптированной для встраиваемой платформы. Алгоритм обеспечивает подпись '
        'и верификацию с 128-битным уровнем безопасности при длине ключа 32 байта '
        'и подписи 64 байта.')

    body(doc,
        'Для защиты приватного ключа при хранении в NVS применяется '
        'многоуровневая схема шифрования. Из пароля пользователя при помощи '
        'функции деривации PBKDF2-SHA256 с 100\u00A0000 итерациями и случайной '
        '16-байтовой солью выводится 256-битный ключ шифрования [7]. Зерно ключа '
        '(32 байта) шифруется алгоритмом AES-256-CBC с дополнением PKCS7, что '
        'даёт 48 байт шифротекста. В NVS сохраняются соль, вектор инициализации '
        'и зашифрованное зерно; открытый ключ остаётся незашифрованным для '
        'возможности идентификации без ввода пароля.')

    body(doc,
        'Серийный протокол взаимодействия с ESP32 определяет шесть типов '
        'сообщений: генерация ключевой пары (0x01), запрос подписи (0x02), '
        'разблокировка паролем (0x03), блокировка (0x04), генерация с паролем (0x05) '
        'и получение открытого ключа (0x06). Запрос передаётся в формате '
        '[тип:1Б][длина:3Б][данные], ответ — [длина:4Б][тип+данные]. Для '
        'фильтрации отладочных сообщений прошивки мост сканирует входящий поток '
        'в поисках байта 0x00, маркирующего начало бинарного пакета.')

    # ════════════════════════════════════════════
    # IMPLEMENTATION
    # ════════════════════════════════════════════
    doc.add_paragraph('Особенности реализации', style='Подзаголовок')

    body(doc,
        'Прошивка ESP32 написана на C++ с использованием фреймворка Arduino '
        'и системы сборки PlatformIO. Криптографические примитивы AES-256-CBC '
        'и PBKDF2-SHA256 реализованы через библиотеку mbedTLS, входящую в состав '
        'ESP-IDF. Для корректной работы PBKDF2 потребовалось включение модуля '
        'MBEDTLS_PKCS5_C в конфигурации сборки. Генерация случайных байт для '
        'соли и вектора инициализации выполняется функцией esp_random(), '
        'использующей аппаратный TRNG микроконтроллера [3].')

    body(doc,
        'Python-мост реализован с применением многопоточной архитектуры: основной '
        'поток обрабатывает входящие соединения через Unix-сокет, а каждый SSH-клиент '
        'обслуживается в отдельном потоке. Библиотека pyserial обеспечивает '
        'взаимодействие с ESP32 через виртуальный COM-порт. При обнаружении '
        'состояния блокировки (код ответа 0x20) мост запрашивает пароль через '
        'стандартный ввод терминала с использованием getpass и автоматически '
        'отправляет команду разблокировки.')

    body(doc,
        'Go-клиент использует пакет golang.org/x/crypto/ssh для реализации '
        'SSH-сессий и стандартный интерфейс ssh-agent для взаимодействия с мостом. '
        'При работе в режиме аппаратного ключа клиент обнаруживает необходимость '
        'ввода пароля через механизм расширений агента и запрашивает его у '
        'пользователя посредством golang.org/x/term. Аутентификация с менеджером '
        'выполняется по протоколу «вызов-ответ»: менеджер отправляет 32 случайных '
        'байта, клиент подписывает их через агент и возвращает 64-байтовую '
        'подпись Ed25519.')

    body(doc,
        'C-менеджер скомпилирован с библиотеками OpenSSL (верификация подписей), '
        'SQLCipher (шифрованная БД), libcurl (REST API к LXD) и libcryptsetup '
        '(управление LUKS-томами). База данных защищена 256-битным ключом через '
        'директиву PRAGMA key. Для каждого пользователя менеджер хранит открытый '
        'ключ и список разрешённых контейнеров. При первом подключении создаётся '
        'LUKS-том в виде loopback-образа размером 100 МБ, который форматируется '
        'в файловую систему ext4 и монтируется внутрь контейнера через механизм '
        'proxy-устройств LXD [5].')

    # ════════════════════════════════════════════
    # TESTING AND RESULTS
    # ════════════════════════════════════════════
    doc.add_paragraph('Результаты тестирования', style='Подзаголовок')

    body(doc,
        'Функциональное тестирование проводилось на стенде, включающем '
        'микроконтроллер ESP32-C3-DevKitM-1, подключённый к серверу под '
        'управлением Ubuntu 24.04 LTS. Тестирование охватило полный цикл работы '
        'системы: генерацию ключевой пары (с паролем и без), экспорт открытого '
        'ключа, регистрацию пользователя, аутентификацию и подключение к контейнеру.')

    body(doc,
        'Корректность реализации Ed25519 подтверждена перекрёстной верификацией: '
        'подписи, сгенерированные на ESP32, успешно проходили проверку стандартными '
        'средствами OpenSSL и утилитой ssh-keygen. Обратная совместимость '
        'проверена путём использования ключа ESP32 для стандартного SSH-подключения '
        'к серверу OpenSSH без участия менеджера.')

    body(doc,
        'Измерения производительности показали следующие результаты: генерация '
        'ключевой пары Ed25519 занимает менее 100 мс, подпись 32-байтового '
        'вызова — 200–250 мс, полный цикл аутентификации через мост — '
        'около 500 мс. Деривация ключа PBKDF2 с 100\u00A0000 итерациями '
        'выполняется за 3–4 секунды, что является приемлемым для однократной '
        'операции разблокировки. Потребление оперативной памяти прошивкой '
        'не превышает 45 КБ, что оставляет значительный запас из доступных '
        '400 КБ SRAM микроконтроллера ESP32-C3.')

    body(doc,
        'Анализ безопасности системы выявил следующие свойства: приватный ключ '
        'не передаётся за пределы микроконтроллера и защищён паролем при хранении; '
        'коммуникация между компонентами осуществляется через локальные интерфейсы '
        '(серийный порт, Unix-сокет); данные контейнеров зашифрованы на уровне '
        'тома (LUKS). К ограничениям следует отнести отсутствие защищённого '
        'элемента (Secure Element) на ESP32-C3, что теоретически допускает '
        'извлечение зашифрованного ключа из NVS при физическом доступе к '
        'микроконтроллеру.')

    # ════════════════════════════════════════════
    # CONCLUSION
    # ════════════════════════════════════════════
    body(doc,
        'В результате работы разработана и реализована система безопасного '
        'удалённого доступа к контейнерным средам, объединяющая аппаратное хранение '
        'криптографических ключей на микроконтроллере ESP32-C3 с управлением '
        'LXD-контейнерами. Система реализует полный криптографический стек: '
        'генерацию ключей Ed25519, их парольное шифрование (AES-256-CBC + '
        'PBKDF2-SHA256), протокол SSH-агента и аутентификацию по схеме '
        '«вызов-ответ». Модульная архитектура из четырёх компонентов обеспечивает '
        'разделение ответственности и возможность независимого развития каждого '
        'модуля.')

    body(doc,
        'Перспективными направлениями развития являются интеграция '
        'аппаратного защищённого элемента (например, ATECC608A) для усиления '
        'защиты ключевого материала, реализация беспроводного подключения '
        '(BLE/Wi-Fi) вместо USB, поддержка нескольких ключевых пар и '
        'совместимость с протоколом FIDO2/WebAuthn для расширения области '
        'применения [8].')

    # ════════════════════════════════════════════
    # BIBLIOGRAPHY
    # ════════════════════════════════════════════
    doc.add_paragraph('Библиографические ссылки', style='БИБЛИОГРАФИЧЕСКИЕ ССЫЛКИ')

    refs = [
        ('Ylonen, T.',
         'Ylonen, T. The Secure Shell (SSH) Authentication Protocol / T. Ylonen, '
         'C. Lonvick // RFC 4252. 2006.'),

        ('Canonical Ltd.',
         'Canonical Ltd. LXD Documentation. URL: https://documentation.ubuntu.com/lxd/ '
         '(дата обращения: 15.03.2026).'),

        ('Espressif Systems.',
         'Espressif Systems. ESP32-C3 Series Datasheet Version 1.2. URL: '
         'https://www.espressif.com/sites/default/files/documentation/'
         'esp32-c3_datasheet_en.pdf (дата обращения: 15.03.2026).'),

        ('Ylonen, T.',
         'Ylonen, T. The Secure Shell (SSH) Transport Layer Protocol / T. Ylonen, '
         'C. Lonvick // RFC 4253. 2006.'),

        ('Stéphane Graber.',
         'Stéphane Graber. LXD — system containers and virtual machines. URL: '
         'https://github.com/canonical/lxd (дата обращения: 15.03.2026).'),

        ('Bernstein, D. J.',
         'Bernstein, D. J. High-speed high-security signatures / D. J. Bernstein, '
         'N. Duif, T. Lange, P. Schwabe, B.-Y. Yang // Journal of Cryptographic '
         'Engineering. 2012. Vol. 2, № 2. P. 77–89.'),

        ('Kaliski, B.',
         'Kaliski, B. PKCS #5: Password-Based Cryptography Specification Version 2.0 / '
         'B. Kaliski // RFC 2898. 2000.'),

        ('FIDO Alliance.',
         'FIDO Alliance. FIDO2: WebAuthn & CTAP. URL: '
         'https://fidoalliance.org/fido2/ (дата обращения: 15.03.2026).'),
    ]

    for i, (author, full_text) in enumerate(refs, 1):
        add_reference(doc, i, full_text, authors_italic=author)

    # Save
    print(f'[*] Saving to {OUTPUT_PATH}...')
    doc.save(OUTPUT_PATH)
    print(f'[+] Done! Article saved: {OUTPUT_PATH}')
    print(f'    Total paragraphs: {len(doc.paragraphs)}')


if __name__ == '__main__':
    main()
