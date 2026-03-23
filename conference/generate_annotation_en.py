#!/usr/bin/env python3
"""
Generate English annotation for CTDA 2024 article.

Output: annotation_en.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
OUTPUT_PATH = os.path.join(PROJECT_ROOT, 'annotation_en.docx')


def setup_styles(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    styles = doc.styles

    if 'Название статьи CTDA' not in [s.name for s in styles]:
        s = styles.add_style('Название статьи CTDA', 1)
    else:
        s = styles['Название статьи CTDA']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(12)
    s.font.bold = True
    s.font.all_caps = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(12)
    s.paragraph_format.line_spacing = 1.0

    if 'Автор' not in [s.name for s in styles]:
        s = styles.add_style('Автор', 1)
    else:
        s = styles['Автор']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(11)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.line_spacing = 1.0

    if 'Аффилиация' not in [s.name for s in styles]:
        s = styles.add_style('Аффилиация', 1)
    else:
        s = styles['Аффилиация']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.italic = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(12)
    s.paragraph_format.line_spacing = 1.0

    if 'Аннотация' not in [s.name for s in styles]:
        s = styles.add_style('Аннотация', 1)
    else:
        s = styles['Аннотация']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.first_line_indent = Cm(1.0)

    if 'Ключевые слова' not in [s.name for s in styles]:
        s = styles.add_style('Ключевые слова', 1)
    else:
        s = styles['Ключевые слова']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(18)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.first_line_indent = Cm(1.0)

    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0


def add_keywords(doc, text):
    p = doc.add_paragraph(style='Ключевые слова')
    run_label = p.add_run('Keywords: ')
    run_label.italic = True
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(10)
    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(10)


def main():
    print('[*] Generating English annotation...')
    doc = Document()
    setup_styles(doc)

    if doc.paragraphs and doc.paragraphs[0].text == '':
        p_el = doc.paragraphs[0]._element
        p_el.getparent().remove(p_el)

    # Title
    doc.add_paragraph(
        'DEVELOPMENT OF A SECURE REMOTE VIRTUALIZATION SYSTEM',
        style='Название статьи CTDA'
    )

    # Author
    doc.add_paragraph('I. I. Ivanov', style='Автор')

    # Affiliation
    p = doc.add_paragraph(style='Аффилиация')
    run = p.add_run(
        'Belarusian State University, Nezavisimosti Av., 4,\n'
        '220030, Minsk, Belarus, ivanov@bsu.by'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph(style='Аффилиация')
    run = p.add_run('Supervisor — P. P. Petrov, Senior Lecturer')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

    # Annotation
    doc.add_paragraph(
        'The architecture and implementation of a system for secure remote access '
        'to containerized environments with hardware-based cryptographic key storage on '
        'an ESP32-C3 microcontroller are presented. The system provides generation, encryption, '
        'and usage of Ed25519 keys directly on the microcontroller, eliminating '
        'the storage of private keys on the user\'s computer. An SSH agent protocol '
        'is implemented via a serial bridge in Python, a Go client for authentication, and '
        'a C manager for LXD container management with LUKS volume encryption. '
        'Testing results confirm the correctness of cryptographic operations '
        'and acceptable signing performance (under 250 ms).',
        style='Аннотация'
    )

    # Keywords
    add_keywords(
        doc,
        'hardware key storage; ESP32; Ed25519; SSH agent; '
        'container virtualization; LXD; LUKS; cryptography.'
    )

    print(f'[*] Saving to {OUTPUT_PATH}...')
    doc.save(OUTPUT_PATH)
    print(f'[+] Done! English annotation saved: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
